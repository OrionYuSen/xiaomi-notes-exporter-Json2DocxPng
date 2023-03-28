import json
import docx
from PIL import Image, ImageDraw, ImageFont
import textwrap


def main():
    # 打开json文件
    with open('notes.json', 'r', encoding='utf-8') as f:
        data = json.load(f)
    # 遍历note
    for note in data['notes']:
        if note['content'] != "":
            # 获得标题和内容
            title = json.loads(note['extraInfo'])['title']  # 由于note['extraInfo']是字符串所以需要json.loads将字符串转换为dictionary
            content = note['content']
            print(title)
            print(content)
            if title != "":  # 空标题会在保存时出错
                # 保存为docx文档
                doc = docx.Document()
                doc.add_paragraph(title)
                doc.add_paragraph(content)
                doc.save(f"{title}.docx")

                # 保存为png
                draw(title, content)


def draw(title, content):
    # 定义图片中文字内容
    content = title + '\n' + content
    # 设置参数
    width = 1920
    bg_color = (255, 255, 255)

    # 定义要绘制的文本和字体
    text = content
    font = ImageFont.truetype('fonts/NotoSerifSC-SemiBold.otf', size=40)

    # 获得文本尺寸并设置行距
    text_width, text_height = font.getsize(text)
    line_height = int(text_height * 1.5)

    # 由于有的笔记很长，所以将笔记分行
    text_lines = textwrap.wrap(text, width=47)

    # 获得文本高度
    text_height = len(text_lines) * line_height

    # 创建图像
    image = Image.new('RGB', (width, text_height), bg_color)
    draw = ImageDraw.Draw(image)

    # 在图片上逐行绘制文本
    y_text = 10
    for line in text_lines:
        draw.text((10, y_text), line, font=font, fill=(0, 0, 0))
        y_text += line_height

    # 在左右两侧增加 20 像素的空隙
    padding = 30
    new_width = width + 2 * padding
    new_image = Image.new('RGB', (new_width, text_height), bg_color)
    new_image.paste(image, (padding, 0))

    # 保存图像
    print(f"{title}.png")
    new_image.save(f"{title}.png")


if __name__ == '__main__':
    main()
